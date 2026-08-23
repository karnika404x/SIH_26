"""
Sovereign AI Workbench - Document Generator
=============================================
Ye file Priyanshi ka kaam hai.

Ye AI ke text output ko ek real Word (.docx) file bana deta hai,
jaise "approval note" jo demo me dikhaya jayega.

Isse alag se run nahi karna - ye main.py ke andar automatically
import ho jata hai.
"""

import os
from fastapi import APIRouter
from pydantic import BaseModel
from docx import Document

router = APIRouter()

# Generated files yahan save honge
OUTPUT_FOLDER = "../demo-files"
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


class DocRequest(BaseModel):
    title: str
    content: str


@router.post("/generate-doc")
def generate_doc(req: DocRequest):
    doc = Document()
    doc.add_heading(req.title, level=0)
    doc.add_paragraph(req.content)

    safe_filename = "".join(c for c in req.title if c.isalnum() or c in (" ", "_", "-")).strip()
    if not safe_filename:
        safe_filename = "approval_note"

    file_path = os.path.join(OUTPUT_FOLDER, f"{safe_filename}.docx")
    doc.save(file_path)

    return {
        "message": "Document generated successfully",
        "file_path": file_path,
    }
